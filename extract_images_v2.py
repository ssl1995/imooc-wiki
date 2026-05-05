from docx import Document
from docx.oxml.ns import qn
import os

doc = Document('word/第0章/审查/5200090_宋声林_融合图像与地理位置元数据的古树名木多模态哈希检索系统研建-盲审.docx')

image_dir = 'word/第0章/审查/extracted_images_v2'
os.makedirs(image_dir, exist_ok=True)

# 建立rel_id到文件名的映射
rel_map = {}
for rel in doc.part.rels.values():
    if 'image' in rel.reltype:
        ext = rel.target_part.content_type.split('/')[-1]
        if ext == 'jpeg':
            ext = 'jpg'
        rel_map[rel.rId] = (rel.target_part.blob, ext)

# 按段落顺序提取图片
img_count = 0
for para_idx, para in enumerate(doc.paragraphs):
    for run in para.runs:
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        for d in drawings:
            embed = d.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed and embed in rel_map:
                img_count += 1
                blob, ext = rel_map[embed]
                # 查找最近的图标题
                title = ""
                for i in range(para_idx, min(para_idx+3, len(doc.paragraphs))):
                    t = doc.paragraphs[i].text.strip()
                    if t.startswith('图') or t.startswith('Figure'):
                        title = t.replace(' ', '_').replace('/', '_')[:50]
                        break
                
                fname = f"para{para_idx}_{title}.{ext}" if title else f"para{para_idx}_img{img_count}.{ext}"
                fpath = os.path.join(image_dir, fname)
                with open(fpath, 'wb') as f:
                    f.write(blob)
                print(f"Saved: {fname} ({len(blob)} bytes)")

print(f"\nTotal: {img_count} images")
