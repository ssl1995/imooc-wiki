from docx import Document

doc = Document('word/第0章/审查/5200090_宋声林_融合图像与地理位置元数据的古树名木多模态哈希检索系统研建-盲审.docx')

# 找到所有图片位置
image_positions = []
for para_idx, para in enumerate(doc.paragraphs):
    for run in para.runs:
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if drawings:
            image_positions.append(para_idx)

# 找到所有图标题
figure_titles = []
for para_idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('图') or text.startswith('Figure'):
        figure_titles.append((para_idx, text))

print("=== 图片位置 ===")
for pos in image_positions:
    print(f"Para {pos}")

print("\n=== 图标题位置 ===")
for idx, text in figure_titles:
    print(f"Para {idx}: {text[:70]}")

print("\n=== 图片与最近图标题匹配 ===")
for img_pos in image_positions:
    # 找最近的图标题（通常在图片后1-2段）
    closest = None
    min_dist = 9999
    for title_pos, title_text in figure_titles:
        dist = abs(title_pos - img_pos)
        if dist < min_dist:
            min_dist = dist
            closest = (title_pos, title_text)
    if closest:
        print(f"Img at para {img_pos} -> Title at para {closest[0]} ({min_dist}段距离): {closest[1][:60]}")
